"""
Extract raw text from uploaded files.
Images: Windows built-in OCR (PowerShell) → best Chinese support, no extra install.
Fallback: pytesseract.
"""
import io
import logging
import subprocess
import tempfile
from pathlib import Path

log = logging.getLogger("file_parser")

# ── Tesseract (fallback only) ─────────────────────────────────────────────────
try:
    import pytesseract
    for _t in [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        "/mnt/c/Program Files/Tesseract-OCR/tesseract.exe",     # WSL path
        "/mnt/c/Program Files (x86)/Tesseract-OCR/tesseract.exe",
    ]:
        if Path(_t).exists():
            pytesseract.pytesseract.tesseract_cmd = _t
            break
    _TESSERACT_OK = True
except ImportError:
    _TESSERACT_OK = False


def extract_text(filename: str, content: bytes) -> str:
    """Return extracted plain text. Error messages are wrapped in [ ]."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _from_pdf(content)
    if ext == ".docx":
        return _from_docx(content)
    if ext == ".doc":
        return "[不支援舊版 .doc 格式，請另存為 .docx 後重新上傳]"
    if ext in (".xlsx", ".xls"):
        return _from_excel(content)
    if ext in (".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"):
        return _from_image(content)
    return f"[不支援的檔案格式：{ext}]"


# ── Text-based formats ────────────────────────────────────────────────────────

def _from_pdf(content: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages)
    except Exception as e:
        return f"[PDF 解析失敗：{e}]"


def _from_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX 解析失敗：{e}]"


def _from_excel(content: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(v) for v in row if v is not None)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[Excel 解析失敗：{e}]"


# ── Image: Windows OCR → tesseract fallback ──────────────────────────────────

def _from_image(content: bytes) -> str:
    """
    Primary: Windows 10/11 built-in OCR via PowerShell (supports zh-Hant natively).
    Fallback: pytesseract with chi_tra+eng.
    Always returns plain text — no sentinel paths.
    """
    text = _windows_ocr(content)
    if text and len(text.strip()) > 3:
        log.info("Windows OCR extracted %d chars", len(text))
        return text

    log.warning("Windows OCR returned empty result, trying tesseract")
    return _tesseract_ocr(content)


def _windows_ocr(content: bytes) -> str:
    """Use Windows Runtime OCR engine via PowerShell (or powershell.exe on WSL)."""
    # Resolve PowerShell binary: prefer native, fall back to WSL interop
    import shutil
    ps_bin = shutil.which("powershell") or shutil.which("powershell.exe") or "powershell.exe"

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    try:
        tmp.write(content)
        tmp.flush()
        tmp.close()

        # On WSL the temp path is a Linux path; convert it to a Windows UNC path
        # so that powershell.exe (Windows process) can open the file.
        wslpath_result = subprocess.run(
            ["wslpath", "-w", tmp.name], capture_output=True, timeout=5
        )
        if wslpath_result.returncode == 0:
            win_path = wslpath_result.stdout.decode().strip()
        else:
            win_path = tmp.name  # already a Windows path (non-WSL)
        img_path = win_path.replace("\\", "\\\\")  # escape for PS string

        ps_script = f"""
[Console]::OutputEncoding = [System.Text.Encoding]::UTF8
Add-Type -AssemblyName System.Runtime.WindowsRuntime
$null = [Windows.Storage.StorageFile,            Windows.Storage,   ContentType=WindowsRuntime]
$null = [Windows.Media.Ocr.OcrEngine,            Windows.Foundation, ContentType=WindowsRuntime]
$null = [Windows.Graphics.Imaging.BitmapDecoder, Windows.Foundation, ContentType=WindowsRuntime]
function Await($t) {{
    $task = [System.WindowsRuntimeSystemExtensions]::AsTask($t)
    $task.Wait(-1) | Out-Null
    return $task.Result
}}
$file   = Await([Windows.Storage.StorageFile]::GetFileFromPathAsync("{img_path}"))
$stream = Await($file.OpenReadAsync())
$dec    = Await([Windows.Graphics.Imaging.BitmapDecoder]::CreateAsync($stream))
$bmp    = Await($dec.GetSoftwareBitmapAsync())
$lang   = [Windows.Globalization.Language]::new("zh-Hant")
$eng    = [Windows.Media.Ocr.OcrEngine]::TryCreateFromLanguage($lang)
if (-not $eng) {{ $eng = [Windows.Media.Ocr.OcrEngine]::TryCreateFromUserProfileLanguages() }}
$res    = Await($eng.RecognizeAsync($bmp))
Write-Output $res.Text
"""
        r = subprocess.run(
            [ps_bin, "-NoProfile", "-Command", ps_script],
            capture_output=True,
            timeout=30,
        )
        if r.returncode == 0:
            return r.stdout.decode("utf-8", errors="replace").strip()
        log.warning("Windows OCR PS exit %d: %s", r.returncode,
                    r.stderr.decode("utf-8", errors="replace")[:200])
        return ""
    except Exception as e:
        log.warning("Windows OCR exception: %s", e)
        return ""
    finally:
        try:
            import os; os.unlink(tmp.name)
        except Exception:
            pass


def _tesseract_ocr(content: bytes) -> str:
    """
    Call tesseract.exe directly via subprocess with WSL path conversion.
    Bypasses pytesseract's temp-file management so Windows binary can read the file.
    """
    # Find tesseract binary
    tesseract_bin = pytesseract.pytesseract.tesseract_cmd if _TESSERACT_OK else None
    if not tesseract_bin or not Path(tesseract_bin).exists():
        return "[圖片解析失敗：Windows OCR 不可用且 Tesseract 未安裝]"

    import os, uuid
    from PIL import Image

    # Upscale small images for better recognition
    try:
        img = Image.open(io.BytesIO(content))
        w, h = img.size
        if w < 1000:
            scale = max(2, 1000 // w)
            img = img.resize((w * scale, h * scale), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        content = buf.getvalue()
    except Exception:
        pass  # use original bytes if PIL fails

    base = f"/tmp/tess_{uuid.uuid4().hex}"
    input_path = f"{base}.png"
    output_base = base

    try:
        with open(input_path, "wb") as f:
            f.write(content)

        # Convert Linux paths to Windows UNC paths so tesseract.exe can access them
        def to_win(path: str) -> str:
            r = subprocess.run(["wslpath", "-w", path], capture_output=True, timeout=5)
            return r.stdout.decode().strip() if r.returncode == 0 else path

        r = subprocess.run(
            [tesseract_bin, to_win(input_path), to_win(output_base), "-l", "chi_tra+eng"],
            capture_output=True,
            timeout=60,
        )
        output_txt = f"{output_base}.txt"
        if os.path.exists(output_txt):
            with open(output_txt, encoding="utf-8") as f:
                text = f.read()
            log.info("Tesseract extracted %d chars", len(text))
            return text
        stderr = r.stderr.decode("utf-8", errors="replace")
        return f"[圖片 OCR 失敗：tesseract exit {r.returncode}: {stderr[:100]}]"
    except Exception as e:
        return f"[圖片 OCR 失敗：{e}]"
    finally:
        for p in [input_path, f"{output_base}.txt"]:
            try:
                os.unlink(p)
            except Exception:
                pass
