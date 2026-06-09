"""
Generate DOCX / PDF snapshots that mirror the company-modal visual style.

Color palette (matches style.css):
  dark-bg      #1e2d4d   header background
  accent       #2e6da8   section titles, table headers
  accent-light #edf4fb   table header fill, section bottom border
  border       #d1dae6   table cell borders
  text         #1a2537   body text
  muted        #7990a8   labels, secondary text
  alt-row      #f7fafd   alternating table row
"""
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Colour tokens ─────────────────────────────────────────────────────────────

class _R:   # python-docx RGBColor constants
    ACCENT       = RGBColor(0x2E, 0x6D, 0xA8)
    ACCENT_LIGHT = RGBColor(0xED, 0xF4, 0xFB)
    TEXT         = RGBColor(0x1A, 0x25, 0x37)
    MUTED        = RGBColor(0x79, 0x90, 0xA8)
    DARK_BG      = RGBColor(0x1E, 0x2D, 0x4D)
    WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
    ALT_ROW      = RGBColor(0xF7, 0xFA, 0xFD)

class _F:   # PyMuPDF float tuples (0–1)
    ACCENT       = (0.180, 0.428, 0.659)
    ACCENT_LIGHT = (0.929, 0.957, 0.984)
    BORDER       = (0.820, 0.855, 0.902)
    TEXT         = (0.102, 0.145, 0.216)
    MUTED        = (0.475, 0.565, 0.659)
    DARK_BG      = (0.118, 0.176, 0.302)
    ALT_ROW      = (0.969, 0.980, 0.992)
    WHITE        = (1.0,   1.0,   1.0  )


# ── Shared markdown helpers ───────────────────────────────────────────────────

def _md_sections(text: str) -> list[tuple[str, str]]:
    """Split markdown into [(heading, body)] pairs. Strips preamble before first ##."""
    if not text:
        return []
    text = re.sub(r"^##\s+.+公司簡介[^\n]*\n+", "", text)
    # Mirror JS renderSummary: drop anything before the first ## heading
    if not text.lstrip().startswith("#"):
        m = re.search(r"^##", text, re.MULTILINE)
        if m:
            text = text[m.start():]
    parts = re.split(r"^(#{1,3}\s+.+)$", text, flags=re.MULTILINE)
    sections: list[tuple[str, str]] = []
    if parts[0].strip():
        sections.append(("", parts[0].strip()))
    i = 1
    while i < len(parts) - 1:
        sections.append((parts[i].lstrip("#").strip(), parts[i + 1].strip()))
        i += 2
    return sections


def _strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    return re.sub(r"\*(.+?)\*", r"\1", text)


def _prose_sentences(text: str) -> list[str]:
    """長段正文依句號拆成數句（短段或單句不拆），對齊前端 _proseParagraphs。"""
    if len(text) < 100:
        return [text]
    parts = [s.strip() for s in re.split(r"(?<=。)", text) if s.strip()]
    return parts if len(parts) >= 2 else [text]


def _display_comp_name(s: str) -> str:
    """顯示用：去掉法定尾綴但保留括號註記（（本案）/（2308）…）。Mirrors _displayCompName."""
    return re.sub(r"(股份有限公司|有限公司)(?=（|$)", "", s or "").strip()


def _strip_comp_suffix(grid: list[list[str]]) -> None:
    """若為競業表（首欄表頭＝公司名稱），就地去掉資料列首欄的法定尾綴。"""
    if grid and grid[0] and grid[0][0].strip() == "公司名稱":
        for row in grid[1:]:
            if row:
                row[0] = _display_comp_name(row[0])


# ── 補充來源 callout — mirrors _supCallout / .sup-callout in app.js + style.css ──
# bg = 低透明度 tint 壓在白底上的近似純色。

_SUP_RE       = re.compile(r"（(簡報|訪談|介紹|筆記)補充")
_SUP_MARKLEN  = 5   # （ + 簡報/訪談/… (2) + 補充 (2)
_SUP_STYLE = {
    "簡報": {"name": "簡報補充", "border": "0d8f7a", "bg": "eef7f6", "label": "0b6e5f"},
    "訪談": {"name": "訪談補充", "border": "7c5cbf", "bg": "f5f2fa", "label": "5b45a0"},
    "介紹": {"name": "介紹補充", "border": "c47d0a", "bg": "faf5eb", "label": "9a6307"},
    "筆記": {"name": "筆記補充", "border": "5b6b7d", "bg": "f2f3f5", "label": "46535f"},
}


def _hex_to_float(h: str) -> tuple[float, float, float]:
    return (int(h[0:2], 16) / 255, int(h[2:4], 16) / 255, int(h[4:6], 16) / 255)


def _sup_span(s: str, idx: int) -> tuple[str, int]:
    """Span of one「（XX補充…）」note at idx → (inner_text, end). Mirrors _supSpan."""
    sep = s[idx + _SUP_MARKLEN] if idx + _SUP_MARKLEN < len(s) else ""
    if sep == "）":
        return s[idx + _SUP_MARKLEN + 1:], len(s)
    depth, j = 0, idx
    while j < len(s):
        if s[j] == "（":
            depth += 1
        elif s[j] == "）":
            depth -= 1
            if depth == 0:
                j += 1
                break
        j += 1
    inner_start = idx + _SUP_MARKLEN + (1 if sep in ("：", ":") else 0)
    return s[inner_start:j - 1], j


def _split_supplements(line: str) -> list[tuple[str, str, str | None]]:
    """[(kind, text, src)] — kind 'text'|'sup'; sup text is marker-stripped inner. Mirrors _splitSupplements."""
    pieces: list[tuple[str, str, str | None]] = []
    i = 0
    while True:
        m = _SUP_RE.search(line, i)
        if not m:
            if i < len(line):
                pieces.append(("text", line[i:], None))
            break
        idx, src = m.start(), m.group(1)
        if idx > i:
            pieces.append(("text", line[i:idx], None))
        inner, end = _sup_span(line, idx)
        pieces.append(("sup", inner, src))
        i = end
    return pieces


def _bullet_sup_inner(raw: str) -> tuple[str, str] | None:
    """If the whole bullet is a supplement note → (inner, src), else None. Mirrors _bulletSupInner."""
    m = re.match(r"^(\*\*)?（(簡報|訪談|介紹|筆記)補充[）：]", raw)
    if not m:
        return None
    src = m.group(2)
    idx = raw.index("（" + src + "補充")
    if idx + _SUP_MARKLEN < len(raw) and raw[idx + _SUP_MARKLEN] == "）":
        return raw[:idx] + raw[idx + _SUP_MARKLEN + 1:], src
    inner, end = _sup_span(raw, idx)
    return raw[:idx] + inner + raw[end:], src


def _split_callout_body(inner: str) -> list[str]:
    """Split a callout body into paragraphs at nested「（XX補充…）」markers (mirrors _splitCalloutBody)."""
    parts: list[str] = []
    i = 0
    while True:
        m = _SUP_RE.search(inner, i)
        if not m:
            t = inner[i:].strip()
            if t:
                parts.append(t)
            break
        t = inner[i:m.start()].strip()
        if t:
            parts.append(t)
        sub, end = _sup_span(inner, m.start())
        s = sub.strip()
        if s:
            parts.append(s)
        i = end
    return parts or [inner.strip()]


def _inline_sup_segments(text: str) -> list[tuple[str, str | None]]:
    """[(text, src_or_None)] keeping the full marker text for inline-coloured spans. Mirrors _wrapSupplements."""
    segs: list[tuple[str, str | None]] = []
    i = 0
    while True:
        m = _SUP_RE.search(text, i)
        if not m:
            if i < len(text):
                segs.append((text[i:], None))
            break
        idx, src = m.start(), m.group(1)
        if idx > i:
            segs.append((text[i:idx], None))
        _, end = _sup_span(text, idx)
        segs.append((text[idx:end], src))
        i = end
    return segs


def _wrap_smart(text: str, size: float, max_w: float) -> list[str]:
    """Width-aware line wrap: CJK chars ≈ 1.0× size, ASCII ≈ 0.55× size."""
    if not text:
        return [""]
    lines, cur, cur_w = [], "", 0.0
    for ch in text:
        ch_w = size * (1.0 if ord(ch) > 0x3000 else 0.55)
        if cur_w + ch_w > max_w and cur:
            lines.append(cur)
            cur, cur_w = "", 0.0
        cur += ch
        cur_w += ch_w
    return (lines + [cur]) if cur else (lines or [""])


def _fmt_capital(n) -> str:
    return f"NT$ {int(n):,} 元" if n else "—"


def _basic_info_rows(company: dict) -> list[tuple[str, str]]:
    """Ordered (label, value) pairs matching the modal's 基本資料 section."""
    par = company.get("par_value")
    no_par = company.get("no_par_value")
    par_str = f"NT$ {par} 元" if par else ("無票面金額" if no_par else "—")

    ts = company.get("total_shares")
    shares_str = f"{int(ts):,} 股（{int(ts) // 1000:,} 張）" if ts else "—"

    rows: list[tuple[str, str]] = [
        ("統一編號",   company.get("tax_id") or "—"),
        ("公司代表人", company.get("representative") or "—"),
        ("資本總額",   _fmt_capital(company.get("authorized_capital"))),
        ("實收資本額", _fmt_capital(company.get("capital"))),
        ("每股金額",   par_str),
        ("股份總數",   shares_str),
        ("公司所在地", company.get("address") or "—"),
        ("設立日期",   company.get("setup_date") or "—"),
        ("產業別",     ", ".join(company.get("industries") or ([company["industry"]] if company.get("industry") else [])) or "—"),
    ]
    if company.get("website"):
        rows.append(("官方網站", company.get("website")))
    return rows


def _dir_table_rows(directors: list[dict]) -> tuple[list[list[str]], bool]:
    """
    Returns (all_rows, has_shares).
    all_rows = [header_row, *data_rows, total_row].
    """
    has_shares = any(d.get("shares") for d in directors)
    header = ["職稱", "姓名", "所代表法人"]
    if has_shares:
        header.append("持股數")
    header.append("持股比例")

    total_shares, total_ratio, has_ratio = 0, 0.0, False
    data: list[list[str]] = []
    for d in directors:
        ratio = d.get("ratio")
        ratio_str = f"{ratio * 100:.2f}%" if ratio is not None else "—"
        if ratio is not None:
            total_ratio += ratio
            has_ratio = True
        sv = d.get("shares")
        if sv:
            total_shares += sv
        row = [d.get("title") or "—", d.get("name") or "—", d.get("representative_of") or "—"]
        if has_shares:
            row.append(f"{int(sv):,}" if sv else "—")
        row.append(ratio_str)
        data.append(row)

    total_row = ["合計", "", ""]
    if has_shares:
        total_row.append(f"{total_shares:,}" if total_shares else "—")
    total_row.append(f"{total_ratio * 100:.2f}%" if has_ratio else "—")

    return [header] + data + [total_row], has_shares


# ── 大股東（公發公司反查）helpers — mirrors _renderShareholderSection in app.js ──

_HOLDER_CATEGORY_LABEL = {
    "subsidiary":       "子公司",
    "associate":        "關聯企業",
    "fvoci_noncurrent": "FVOCI股權投資（非流動）",
    "fvoci_current":    "FVOCI股權投資（流動）",
    "fvoci_equity":     "FVOCI股權投資",
    "mainland_china":   "大陸投資",
    "other_lt_equity":  "其他長期股權",
}


def _holder_pct_display(r: dict, total_shares: int) -> str:
    """Mirror app.js pctDisplay: 持股張數 + (持股比例)."""
    shares = r.get("shares_nt")
    shares_num = None
    if shares is not None:
        try:
            shares_num = float(shares)
        except (TypeError, ValueError):
            shares_num = None
    shares_str = f"{int(shares_num):,}張" if shares_num is not None else None

    pct = r.get("pct")
    ratio = pct if pct is not None else (
        shares_num * 1000 / total_shares
        if (shares_num is not None and total_shares) else None
    )
    ratio_str = f"({ratio * 100:.2f}%)" if ratio is not None else None

    if shares_str and ratio_str:
        return f"{shares_str} {ratio_str}"
    return shares_str or ratio_str or "—"


def _holder_table_rows(results: list[dict], total_shares: int) -> list[list[str]]:
    header = ["持有公司", "代號", "持股張數/比例", "資料日期", "類型"]
    rows = [header]
    for r in results:
        cat = r.get("category") or ""
        rows.append([
            r.get("holder_name") or "—",
            r.get("holder_id") or "—",
            _holder_pct_display(r, total_shares),
            r.get("as_of_date") or "—",
            _HOLDER_CATEGORY_LABEL.get(cat, cat or "—"),
        ])
    return rows


def _shareholder_block(company: dict, holders: dict | None) -> dict | None:
    """
    Return a render plan for the 大股東 section, or None to hide it.
    Mirrors _renderShareholderSection: only shown when directors carry ratios.
    """
    directors = company.get("directors") or []
    ratios = [d["ratio"] for d in directors if d.get("ratio") is not None]
    if not ratios:
        return None
    total_ratio = sum(ratios)
    pct = total_ratio * 100
    if total_ratio >= 0.999:
        return {"complete": True,
                "note": f"董監事持股合計 {pct:.2f}%，持股已完整揭露"}

    missing = 100 - pct
    block: dict = {
        "complete": False,
        "alert": (f"董監事持股合計 {pct:.2f}%，尚有 {missing:.2f}% "
                  f"股份未在董監事名單中揭露，可能由其他股東持有"),
    }
    results = (holders or {}).get("results") or []
    if results:
        block["found_note"] = f"找到 {len(results)} 家公發公司揭露持有此公司股份："
        block["rows"] = _holder_table_rows(results, (holders or {}).get("total_shares") or 0)
    else:
        block["empty_note"] = "查無公發公司揭露持有此公司股份"
    return block


def _patent_table_rows(patents: list[dict]) -> list[list[str]]:
    header = ["專利號", "名稱", "申請日", "狀態", "發明人"]
    rows = [header]
    for p in patents:
        rows.append([
            p.get("patent_no") or "—",
            p.get("title") or "—",
            p.get("app_date") or "—",
            p.get("status") or "—",
            "、".join(p.get("inventors") or []) or "—",
        ])
    return rows


# ── DOCX width helpers ────────────────────────────────────────────────────────

def _docx_est_pt(text: str, size_pt: float = 9) -> float:
    """Estimate text width in points (CJK ≈ 1.0×, ASCII ≈ 0.55×)."""
    return sum(size_pt * (1.0 if ord(ch) > 0x00FF else 0.55)
               for ch in _strip_inline_md(text))


def _docx_auto_cols(rows: list[list[str]], size_pt: float,
                    total_pt: float, pad_pt: float = 18) -> list:
    """Return Pt() column widths fitted to content, summing to total_pt."""
    if not rows:
        return []
    n = max(len(r) for r in rows)
    ideal = []
    for ci in range(n):
        mx = max((_docx_est_pt(str(r[ci]), size_pt) if ci < len(r) else 0)
                 for r in rows) + pad_pt
        ideal.append(mx)
    total = sum(ideal)
    if total <= total_pt:
        extra = total_pt - total
        if n > 2:
            mid = max(range(1, n - 1), key=lambda i: ideal[i])
            ideal[mid] += extra
    else:
        scale = total_pt / total
        ideal = [w * scale for w in ideal]
    return [Pt(w) for w in ideal]


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX
# ══════════════════════════════════════════════════════════════════════════════

# ── DOCX XML helpers ──────────────────────────────────────────────────────────

def _shading(cell, hex_fill: str, hex_color: str = "auto"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), hex_color)
    shd.set(qn("w:fill"),  hex_fill)
    tcPr.append(shd)


def _cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _no_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBdr = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBdr.append(el)
    tblPr.append(tblBdr)


def _set_cell_borders(cell, hex_color: str = "d1dae6", sides=("bottom",), sz: int = 4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        tcBdr.append(el)
    tcPr.append(tcBdr)


def _para_bottom_border(para, hex_color: str = "c7d9ee", sz: int = 4):
    pPr   = para._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _para_left_border(para, hex_color: str = "2e6da8", sz: int = 18):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(sz))
    left.set(qn("w:space"), "9")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "180")
    pPr.append(ind)


def _para_spacing(para, before: int = 0, after: int = 0, line: int | None = None):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"),  str(after))
    if line:
        sp.set(qn("w:line"),     str(line))
        sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _letter_spacing(run, val: int = 40):
    rPr = run._r.get_or_add_rPr()
    sp  = OxmlElement("w:spacing")
    sp.set(qn("w:val"), str(val))
    rPr.append(sp)


def _add_md_runs(para, text: str, size_pt: float, color: RGBColor):
    """Parse **bold** inline markdown and add styled runs."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        bold = part.startswith("**") and part.endswith("**")
        r = para.add_run(part[2:-2] if bold else part)
        r.bold = bold
        r.font.size  = Pt(size_pt)
        r.font.color.rgb = color


def _add_rich_runs(para, text: str, size_pt: float, base_color: RGBColor):
    """Like _add_md_runs but colours inline 「（XX補充…）」spans by source (mirrors _wrapSupplements)."""
    for seg, src in _inline_sup_segments(text):
        color = RGBColor.from_string(_SUP_STYLE[src]["label"]) if src else base_color
        for part in re.split(r"(\*\*[^*]+\*\*)", seg):
            if not part:
                continue
            bold = part.startswith("**") and part.endswith("**")
            r = para.add_run(part[2:-2] if bold else part)
            r.bold = bold
            r.font.size = Pt(size_pt)
            r.font.color.rgb = color


# ── DOCX builder helpers ──────────────────────────────────────────────────────

def _docx_section_heading(doc, text: str):
    """Section title: uppercase, accent blue, bottom border — mirrors .modal-section h4"""
    p = doc.add_paragraph()
    _para_spacing(p, before=200, after=80)
    _para_bottom_border(p, "c7d9ee", sz=4)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = _R.ACCENT
    _letter_spacing(r, 40)
    return p


def _docx_summary_h3(doc, text: str):
    """Summary ## heading: left accent border, mirrors #modal-summary h3"""
    p = doc.add_paragraph()
    _para_spacing(p, before=180, after=60)
    _para_left_border(p, "2e6da8", sz=18)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = _R.TEXT
    return p


def _docx_summary_h4(doc, text: str):
    p = doc.add_paragraph()
    _para_spacing(p, before=80, after=20)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = _R.MUTED
    return p


def _docx_callout(doc, inner: str, src: str):
    """Source-coloured supplement box: tinted fill + left accent bar + label. Mirrors .sup-callout."""
    style = _SUP_STYLE.get(src, _SUP_STYLE["簡報"])
    tbl  = doc.add_table(rows=1, cols=1)
    _no_table_borders(tbl)
    cell = tbl.cell(0, 0)
    _shading(cell, style["bg"])
    _set_cell_borders(cell, style["border"], ("left",), sz=18)   # ~3px accent bar
    _cell_margins(cell, top=70, bottom=70, left=150, right=130)

    lp = cell.paragraphs[0]
    _para_spacing(lp, before=0, after=40)
    lr = lp.add_run(style["name"])
    lr.bold = True
    lr.font.size = Pt(8.5)
    lr.font.color.rgb = RGBColor.from_string(style["label"])

    parts = _split_callout_body(inner)
    for pi, part in enumerate(parts):
        bp = cell.add_paragraph()
        _para_spacing(bp, before=0, after=(60 if pi < len(parts) - 1 else 0), line=270)
        _add_rich_runs(bp, part, 9, _R.TEXT)

    sp = doc.add_paragraph()
    _para_spacing(sp, before=0, after=0)


def _docx_summary_body(doc, lines: list[str],
                       content_pt: float = 468):
    """Render body lines: lists, tables, paragraphs."""
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip blank lines and horizontal rules
        if not stripped or re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # 競業類型定義行：前端改用頁籤 tooltip，匯出也一併略過
        if re.match(r"^競業類型定義[：:]", stripped):
            i += 1
            continue

        # ### sub-heading
        if stripped.startswith("###"):
            _docx_summary_h4(doc, stripped.lstrip("#").strip())
            i += 1
            continue

        # Markdown table block
        if stripped.startswith("|") and i + 1 < len(lines) and \
                re.match(r"^\s*\|[-| ]+\|\s*$", lines[i + 1]):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            data_rows = [r for r in table_lines
                         if not re.match(r"^\s*\|[-| ]+\|\s*$", r.strip())]
            if data_rows:
                grid = [[c.strip() for c in r.strip().strip("|").split("|")]
                        for r in data_rows]
                _strip_comp_suffix(grid)
                cols  = max(len(r) for r in grid)
                cws   = _docx_auto_cols(grid, 9, content_pt)
                t     = doc.add_table(rows=len(grid), cols=cols)
                t.style = "Table Grid"
                for ci, cw in enumerate(cws):
                    for cell in (t.columns[ci].cells if ci < cols else []):
                        cell.width = cw
                for ri, row_cells in enumerate(grid):
                    for ci, val in enumerate(row_cells[:cols]):
                        cell = t.cell(ri, ci)
                        cell.width = cws[ci] if ci < len(cws) else Pt(content_pt / cols)
                        cell.paragraphs[0].clear()
                        _cell_margins(cell, top=40, bottom=40, left=80, right=80)
                        r = cell.paragraphs[0].add_run(_strip_inline_md(val))
                        r.font.size = Pt(9)
                        if ri == 0:
                            _shading(cell, "edf4fb")
                            r.bold = True
                            r.font.color.rgb = _R.ACCENT
                        else:
                            r.font.color.rgb = _R.TEXT
                            if ri % 2 == 0:
                                _shading(cell, "f7fafd")
            continue

        # List item — hanging indent matching PDF (bullet at 10pt, text at 22pt)
        if re.match(r"^[-*]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
            content = re.sub(r"^[-*]\s+", "", re.sub(r"^\d+\.\s+", "", stripped))
            # Whole-bullet supplement → callout box (mirrors _bulletSupInner path)
            bsup = _bullet_sup_inner(content)
            if bsup:
                _docx_callout(doc, bsup[0], bsup[1])
                i += 1
                continue
            p   = doc.add_paragraph()
            _para_spacing(p, before=0, after=40, line=276)
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"),    str(22 * 20))   # 22pt in twips
            ind.set(qn("w:hanging"), str(12 * 20))   # 12pt hanging → bullet at 10pt
            pPr.append(ind)
            br = p.add_run("• ")
            br.font.size      = Pt(8)
            br.font.color.rgb = _R.TEXT
            _add_rich_runs(p, content, 9.5, _R.TEXT)   # inline supplements coloured
            i += 1
            continue

        # Regular paragraph — split out supplement notes into callout boxes; 長段正文再依句拆段
        for kind, text, src in _split_supplements(stripped):
            if kind == "sup":
                _docx_callout(doc, text, src)
            elif text.strip():
                for sent in _prose_sentences(text.strip()):
                    p = doc.add_paragraph()
                    _para_spacing(p, before=20, after=60, line=276)
                    _add_md_runs(p, sent, 9.5, _R.TEXT)
        i += 1


def _docx_render_table(doc, grid: list[list[str]], content_pt: float = 468):
    """Render a list-of-rows as a styled table (header shaded, alt rows)."""
    if not grid:
        return
    cols = max(len(r) for r in grid)
    cws  = _docx_auto_cols(grid, 9, content_pt)
    t    = doc.add_table(rows=len(grid), cols=cols)
    t.style = "Table Grid"
    for ri, row_cells in enumerate(grid):
        for ci in range(cols):
            val  = row_cells[ci] if ci < len(row_cells) else ""
            cell = t.cell(ri, ci)
            cell.width = cws[ci] if ci < len(cws) else Pt(content_pt / cols)
            cell.paragraphs[0].clear()
            _cell_margins(cell, top=40, bottom=40, left=80, right=80)
            r = cell.paragraphs[0].add_run(_strip_inline_md(str(val)))
            r.font.size = Pt(9)
            if ri == 0:
                _shading(cell, "edf4fb")
                r.bold = True
                r.font.color.rgb = _R.ACCENT
            else:
                r.font.color.rgb = _R.TEXT
                if ri % 2 == 0:
                    _shading(cell, "f7fafd")


def _docx_note_paragraph(doc, text: str, color, size_pt: float = 9,
                         before: int = 0, after: int = 30):
    p = doc.add_paragraph()
    _para_spacing(p, before=before, after=after, line=276)
    r = p.add_run(text)
    r.font.size = Pt(size_pt)
    r.font.color.rgb = color
    return p


# ── Main DOCX builder ─────────────────────────────────────────────────────────

def build_docx(company: dict, holders: dict | None = None) -> bytes:
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    # ── Dark header block (mirrors #modal-header background: var(--sb-bg)) ──
    header_tbl = doc.add_table(rows=1, cols=1)
    _no_table_borders(header_tbl)
    hcell = header_tbl.cell(0, 0)
    _shading(hcell, "1e2d4d")
    _cell_margins(hcell, top=220, bottom=180, left=300, right=300)

    name_para = hcell.paragraphs[0]
    _para_spacing(name_para, before=0, after=80)
    name_run = name_para.add_run(company.get("name", "—"))
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = _R.WHITE

    listing = company.get("listing_status", "")
    if listing:
        badge_para = hcell.add_paragraph()
        _para_spacing(badge_para, before=0, after=0)
        br = badge_para.add_run(f"[{listing}]")
        br.font.size = Pt(9)
        br.font.color.rgb = _R.MUTED

    labels = company.get("labels") or []
    if labels:
        lp = hcell.add_paragraph()
        _para_spacing(lp, before=60, after=0)
        lr = lp.add_run("  ".join(labels[:8]))
        lr.font.size = Pt(8.5)
        lr.font.color.rgb = RGBColor(0xBB, 0xCC, 0xDD)

    doc.add_paragraph()  # spacer

    # ── Basic info ──
    _docx_section_heading(doc, "基本資料")

    info_tbl = doc.add_table(rows=0, cols=2)
    _no_table_borders(info_tbl)
    info_tbl.columns[0].width = Inches(1.4)
    info_tbl.columns[1].width = Inches(4.2)

    for label, value in _basic_info_rows(company):
        row = info_tbl.add_row()
        lc, vc = row.cells
        _cell_margins(lc, top=30, bottom=30, left=0, right=60)
        _cell_margins(vc, top=30, bottom=30, left=0, right=0)
        _set_cell_borders(lc, "d1dae6", ("bottom",))
        _set_cell_borders(vc, "d1dae6", ("bottom",))
        _para_spacing(lc.paragraphs[0], before=0, after=0)
        lr = lc.paragraphs[0].add_run(label)
        lr.font.size = Pt(9)
        lr.font.color.rgb = _R.MUTED
        _para_spacing(vc.paragraphs[0], before=0, after=0)
        vr = vc.paragraphs[0].add_run(value)
        vr.font.size = Pt(9)
        vr.bold = True
        vr.font.color.rgb = _R.TEXT

    doc.add_paragraph()

    # ── Directors ──
    directors = company.get("directors") or []
    if directors:
        _docx_section_heading(doc, "董監事名單")

        CONTENT_PT = 468  # 6.5 in × 72 pt
        dir_rows, _ = _dir_table_rows(directors)
        n_cols = len(dir_rows[0])
        cws = _docx_auto_cols(dir_rows[:-1], 9, CONTENT_PT)  # exclude total row from width calc

        dt = doc.add_table(rows=1, cols=n_cols)
        dt.style = "Table Grid"
        hdr_row = dt.rows[0]
        for cell, txt_val, cw in zip(hdr_row.cells, dir_rows[0], cws):
            cell.width = cw
            _shading(cell, "edf4fb")
            _cell_margins(cell, top=50, bottom=50, left=80, right=80)
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(txt_val)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = _R.ACCENT

        for idx, vals in enumerate(dir_rows[1:]):
            is_total = idx == len(dir_rows) - 2
            row = dt.add_row()
            if is_total:
                for cell in row.cells:
                    _shading(cell, "edf4fb")
            elif idx % 2 == 1:
                for cell in row.cells:
                    _shading(cell, "f7fafd")
            for cell, val, cw in zip(row.cells, vals, cws):
                cell.width = cw
                _cell_margins(cell, top=40, bottom=40, left=80, right=80)
                cell.paragraphs[0].clear()
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(9)
                r.font.color.rgb = _R.ACCENT if is_total else _R.TEXT
                r.bold = is_total

        doc.add_paragraph()

    # ── 大股東 ──
    sh = _shareholder_block(company, holders)
    if sh:
        _docx_section_heading(doc, "大股東")
        if sh.get("complete"):
            _docx_note_paragraph(doc, sh["note"], _R.MUTED, after=20)
        else:
            _docx_note_paragraph(doc, "⚠ " + sh["alert"], _R.TEXT, after=30)
            if sh.get("rows"):
                _docx_note_paragraph(doc, sh["found_note"], _R.MUTED, after=20)
                _docx_render_table(doc, sh["rows"])
            elif sh.get("empty_note"):
                _docx_note_paragraph(doc, sh["empty_note"], _R.MUTED, after=20)
        doc.add_paragraph()

    # ── Summary ──
    summary_raw = company.get("summary", "")
    if summary_raw:
        _docx_section_heading(doc, "公司簡介")

        for heading, body in _md_sections(summary_raw):
            if heading:
                _docx_summary_h3(doc, heading)
            _docx_summary_body(doc, body.split("\n"))

    # ── 專利 ──
    patents = company.get("patents") or []
    if patents:
        _docx_section_heading(doc, "專利")
        _docx_note_paragraph(
            doc, f"共 {len(patents)} 筆（更新：{patents[0].get('fetched_at', '')}）",
            _R.MUTED, after=20)
        _docx_render_table(doc, _patent_table_rows(patents))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
#  PDF
# ══════════════════════════════════════════════════════════════════════════════

def build_pdf(company: dict, holders: dict | None = None) -> bytes:
    import fitz

    PAGE_W, PAGE_H = 595, 842
    ML, MR = 40, 40
    MB = 65
    CW = PAGE_W - ML - MR    # 475 pt

    doc  = fitz.open()
    page: fitz.Page = None    # type: ignore[assignment]
    y    = 0.0

    # Pre-create font objects for width calculation
    _fhelv = fitz.Font("helv")
    _fcjk  = fitz.Font("china-t")

    # ── Page management ───────────────────────────────────────────────────────

    def _new_page() -> fitz.Page:
        nonlocal page, y
        page = doc.new_page(width=PAGE_W, height=PAGE_H)
        y    = 60.0
        return page

    def _need_space(h: float):
        if y + h > PAGE_H - MB:
            _new_page()

    # ── Drawing primitives ────────────────────────────────────────────────────

    def filled_rect(x0, y0, x1, y1, fill):
        page.draw_rect(fitz.Rect(x0, y0, x1, y1), color=None, fill=fill, width=0)

    def hline(y_pos: float, x0=ML, x1=None, color=_F.BORDER, w=0.5):
        page.draw_line(fitz.Point(x0, y_pos),
                       fitz.Point(x1 or PAGE_W - MR, y_pos),
                       color=color, width=w)

    def vline(x_pos: float, y0: float, y1: float, color=_F.ACCENT, w=3.0):
        page.draw_line(fitz.Point(x_pos, y0),
                       fitz.Point(x_pos, y1),
                       color=color, width=w)

    # ── Text helpers ──────────────────────────────────────────────────────────

    def _seg_width(s: str, size: float, is_cjk: bool) -> float:
        """Estimate segment width using font metrics."""
        if is_cjk:
            return _fcjk.text_length(s, fontsize=size)
        return _fhelv.text_length(s, fontsize=size)

    def _text_segments(s: str) -> list[tuple[str, bool]]:
        """Split into [(text, is_cjk)] alternating runs.
        Threshold >0x00FF: numbers/basic-latin → helv; everything else → china-t.
        This ensures em-dash (U+2014), Chinese, etc. all use the CJK font."""
        segs, cur, cur_cjk = [], "", None
        for ch in s:
            is_cjk = ord(ch) > 0x00FF
            if cur_cjk is not None and is_cjk != cur_cjk:
                segs.append((cur, cur_cjk)); cur = ""
            cur_cjk = is_cjk; cur += ch
        if cur: segs.append((cur, cur_cjk if cur_cjk is not None else False))
        return segs

    def txt(s: str, x: float, y_baseline: float, size: float = 10, color=_F.TEXT):
        """Render text with helv for ASCII/numbers, china-t for CJK."""
        for seg, is_cjk in _text_segments(s):
            fn = "china-t" if is_cjk else "helv"
            page.insert_text(fitz.Point(x, y_baseline), seg,
                             fontname=fn, fontsize=size, color=color)
            x += _seg_width(seg, size, is_cjk)

    def _line_width(s: str, size: float) -> float:
        return sum(_seg_width(seg, size, is_cjk) for seg, is_cjk in _text_segments(s))

    def _wrap_mixed(s: str, size: float, max_w: float) -> list[str]:
        """Width-accurate wrap using font metrics."""
        if not s:
            return [""]
        lines, cur, cur_w = [], "", 0.0
        for ch in s:
            is_cjk = ord(ch) > 0x00FF
            ch_w   = _fcjk.text_length(ch, fontsize=size) if is_cjk \
                     else _fhelv.text_length(ch, fontsize=size)
            if cur_w + ch_w > max_w and cur:
                lines.append(cur); cur, cur_w = "", 0.0
            cur += ch; cur_w += ch_w
        return (lines + [cur]) if cur else (lines or [""])

    def _auto_col_widths(rows: list[list[str]], size: float,
                         total_w: float, pad: float = 10) -> list[float]:
        """Compute column widths so every cell fits in one line, scaled to total_w."""
        if not rows:
            return []
        n = max(len(r) for r in rows)
        ideal = []
        for ci in range(n):
            max_cell = 0.0
            for row in rows:
                if ci < len(row):
                    w = _line_width(_strip_inline_md(str(row[ci])), size)
                    max_cell = max(max_cell, w)
            ideal.append(max_cell + pad)
        total_ideal = sum(ideal)
        if total_ideal <= total_w:
            # Give leftover to the widest middle column
            extra = total_w - total_ideal
            if n > 2:
                mid = max(range(1, n - 1), key=lambda i: ideal[i])
                ideal[mid] += extra
        else:
            scale = total_w / total_ideal
            ideal = [w * scale for w in ideal]
        return ideal

    def put(s: str, size: float = 10, x: float = ML,
            color=_F.TEXT, gap_after: float = 3, max_w: float | None = None) -> None:
        """Wrap and place text, advance y, handle page breaks."""
        nonlocal y
        eff_w = (max_w if max_w is not None else CW) - (x - ML)
        for ln in _wrap_mixed(_strip_inline_md(s), size, eff_w):
            _need_space(size + 4)
            txt(ln, x, y + size, size, color)
            y += size + 3
        y += gap_after

    # ── Section heading ───────────────────────────────────────────────────────

    def section_heading(title: str, gap_before: float = 18):
        nonlocal y
        _need_space(gap_before + 14)
        y += gap_before
        txt(title.upper(), ML, y + 9, size=9, color=_F.ACCENT)
        y += 9 + 4
        hline(y, color=_F.ACCENT_LIGHT, w=2)
        y += 8

    # ── Summary heading helpers ───────────────────────────────────────────────

    def summary_h3(title: str):
        nonlocal y
        _need_space(26)
        y += 12
        vline(ML, y, y + 13, color=_F.ACCENT, w=3)
        # h3 might be long — wrap it within content width minus indent
        for i, ln in enumerate(_wrap_mixed(title, 10, CW - 9)):
            txt(ln, ML + 9, y + 10 + i * 13, size=10, color=_F.TEXT)
        y += 13 + 5

    def summary_h4(title: str):
        nonlocal y
        y += 6
        put(title, size=9, color=_F.MUTED, gap_after=3)

    # ── Table renderer ────────────────────────────────────────────────────────

    def pdf_table(rows: list[list[str]], col_widths: list[float],
                  x0: float = ML, header: bool = True):
        nonlocal y
        n_cols    = len(col_widths)
        total_w   = sum(col_widths)
        CELL_PAD  = 5
        MIN_ROW_H = 18.0

        for ri, row in enumerate(rows):
            is_header = header and ri == 0
            is_alt    = (not is_header) and ri % 2 == 0
            cell_size = 8.5

            # Pre-wrap each cell to know row height
            wrapped_cells = []
            for ci, val in enumerate(row[:n_cols]):
                max_cw = col_widths[ci] - CELL_PAD * 2
                wrapped_cells.append(_wrap_mixed(_strip_inline_md(str(val)), cell_size, max_cw))
            row_lines = max(len(wc) for wc in wrapped_cells)
            row_h     = max(MIN_ROW_H, row_lines * (cell_size + 2) + CELL_PAD * 2)

            _need_space(row_h + 2)

            if is_header:
                filled_rect(x0, y, x0 + total_w, y + row_h, _F.ACCENT_LIGHT)
            elif is_alt:
                filled_rect(x0, y, x0 + total_w, y + row_h, _F.ALT_ROW)

            cx = x0
            for ci, wlines in enumerate(wrapped_cells):
                cell_color = _F.ACCENT if is_header else _F.TEXT
                for li, ln in enumerate(wlines):
                    txt(ln, cx + CELL_PAD, y + CELL_PAD + cell_size + li * (cell_size + 2),
                        size=cell_size, color=cell_color)
                cx += col_widths[ci]

            hline(y + row_h, x0=x0, x1=x0 + total_w, color=_F.BORDER, w=0.4)
            y += row_h
        y += 6

    # ── Supplement callout (source-coloured box) ──────────────────────────────

    def callout(inner: str, src: str):
        """Tinted box + left accent bar + coloured label, body split into 段. Mirrors .sup-callout."""
        nonlocal y
        style    = _SUP_STYLE.get(src, _SUP_STYLE["簡報"])
        bg_c     = _hex_to_float(style["bg"])
        bar_c    = _hex_to_float(style["border"])
        label_c  = _hex_to_float(style["label"])
        PAD, BAR, PARA_GAP = 8.0, 3.0, 5.0
        inner_x  = ML + BAR + PAD
        text_w   = CW - BAR - PAD * 2
        wrapped  = [_wrap_mixed(_strip_inline_md(p), 9, text_w)
                    for p in _split_callout_body(inner)]
        body_h   = (sum(len(w) * (9 + 3) for w in wrapped)
                    + PARA_GAP * (len(wrapped) - 1))
        box_h    = PAD + 10 + 5 + body_h + PAD

        # If the whole box won't fit, start a fresh page so the fill stays intact.
        if y + box_h > PAGE_H - MB and box_h <= PAGE_H - 60 - MB:
            _new_page()

        filled_rect(ML, y, ML + CW, y + box_h, bg_c)      # tinted background
        filled_rect(ML, y, ML + BAR, y + box_h, bar_c)    # left accent bar
        yy = y + PAD
        txt(style["name"], inner_x, yy + 8.5, size=8.5, color=label_c)
        yy += 10 + 5
        for wi, w in enumerate(wrapped):
            for ln in w:
                txt(ln, inner_x, yy + 9, size=9, color=_F.TEXT)
                yy += 9 + 3
            if wi < len(wrapped) - 1:
                yy += PARA_GAP
        y += box_h + 6

    # ── PDF summary body ──────────────────────────────────────────────────────

    def pdf_summary_body(lines: list[str]):
        nonlocal y
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line or re.match(r"^-{3,}$", line):   # skip blank / hr
                y += 4; i += 1; continue

            # 競業類型定義行：前端改用頁籤 tooltip，匯出也一併略過
            if re.match(r"^競業類型定義[：:]", line):
                i += 1; continue

            # Markdown table
            if line.startswith("|") and i + 1 < len(lines) and \
               re.match(r"^\s*\|[-| ]+\|\s*$", lines[i + 1].strip()):
                block = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    block.append(lines[i]); i += 1
                data = [r for r in block if not re.match(r"^\s*\|[-| ]+\|\s*$", r.strip())]
                rows = [[c.strip() for c in r.strip().strip("|").split("|")] for r in data]
                _strip_comp_suffix(rows)
                if rows:
                    n  = max(len(r) for r in rows)
                    pdf_table(rows, _auto_col_widths(rows, 8.5, CW))
                continue

            # List item — hanging indent: bullet at BX, text+wrap at TX
            if re.match(r"^[-*]\s+", line) or re.match(r"^\d+\.\s+", line):
                content = re.sub(r"^[-*]\s+", "", re.sub(r"^\d+\.\s+", "", line))
                # Whole-bullet supplement → callout box (mirrors _bulletSupInner path)
                bsup = _bullet_sup_inner(content)
                if bsup:
                    callout(bsup[0], bsup[1])
                    i += 1; continue
                BX, TX = ML + 10, ML + 22          # bullet pos, text pos
                TW     = CW - (TX - ML)             # wrap width from TX to right margin
                lns    = _wrap_mixed(_strip_inline_md(content), 9.5, TW)
                for li, ln in enumerate(lns):
                    _need_space(9.5 + 3)
                    if li == 0:
                        txt("•", BX, y + 9, size=7, color=_F.TEXT)
                    txt(ln, TX, y + 9.5, size=9.5, color=_F.TEXT)
                    y += 9.5 + 3
                y += 3
                i += 1; continue

            # ### sub-heading
            if line.startswith("###"):
                summary_h4(line.lstrip("#").strip())
                i += 1; continue

            # Regular paragraph — supplement notes → callout boxes; 長段正文再依句拆段
            for kind, text, src in _split_supplements(line):
                if kind == "sup":
                    callout(text, src)
                elif text.strip():
                    for sent in _prose_sentences(text.strip()):
                        put(sent, size=9.5, gap_after=5)
            i += 1

    # ═══════════════ Build document ═══════════════

    _new_page()

    # ── Header block (dark navy) ──────────────────────────────────────────────
    HDR_H = 88
    filled_rect(0, 0, PAGE_W, HDR_H, _F.DARK_BG)

    name       = company.get("name", "—")
    name_lines = _wrap_mixed(name, 18, CW)[:2]
    txt(name_lines[0], ML, 32, size=18, color=_F.WHITE)
    if len(name_lines) > 1:
        txt(name_lines[1], ML, 52, size=18, color=_F.WHITE)

    tags = []
    if company.get("listing_status"):
        tags.append(f"[{company['listing_status']}]")
    tags += (company.get("labels") or [])[:5]
    if tags:
        txt("  ".join(tags), ML, 76, size=8.5, color=_F.MUTED)

    y = HDR_H + 16

    # ── Basic info ────────────────────────────────────────────────────────────
    section_heading("基本資料", gap_before=0)

    COL1     = 88
    VAL_MAXW = CW - COL1

    for label, value in _basic_info_rows(company):
        val_lines = _wrap_mixed(value, 9, VAL_MAXW)
        row_h     = max(16, len(val_lines) * 12 + 4)
        _need_space(row_h)
        txt(label, ML, y + 10, size=9, color=_F.MUTED)
        for li, vl in enumerate(val_lines):
            txt(vl, ML + COL1, y + 10 + li * 12, size=9, color=_F.TEXT)
        hline(y + row_h - 2, color=_F.BORDER, w=0.3)
        y += row_h

    # ── Directors ─────────────────────────────────────────────────────────────
    directors = company.get("directors") or []
    if directors:
        section_heading("董監事名單")
        dir_rows, _ = _dir_table_rows(directors)
        col_widths = _auto_col_widths(dir_rows[:-1], 8.5, CW)  # exclude total from width calc
        pdf_table(dir_rows[:-1], col_widths)

        # Total row — accent-light background, bold accent text
        total = dir_rows[-1]
        n_cols = len(col_widths)
        total_h = 20.0
        total_w = sum(col_widths)
        _need_space(total_h)
        filled_rect(ML, y, ML + total_w, y + total_h, _F.ACCENT_LIGHT)
        cx = ML
        for ci, val in enumerate(total[:n_cols]):
            txt(val, cx + 5, y + total_h - 6, size=8.5, color=_F.ACCENT)
            cx += col_widths[ci]
        hline(y + total_h, x0=ML, x1=ML + total_w, color=_F.BORDER, w=0.4)
        y += total_h + 6

    # ── 大股東 ─────────────────────────────────────────────────────────────────
    sh = _shareholder_block(company, holders)
    if sh:
        section_heading("大股東")
        if sh.get("complete"):
            put(sh["note"], size=9, color=_F.MUTED, gap_after=4)
        else:
            put("⚠ " + sh["alert"], size=9, color=_F.TEXT, gap_after=5)
            if sh.get("rows"):
                put(sh["found_note"], size=9, color=_F.MUTED, gap_after=3)
                rows = sh["rows"]
                pdf_table(rows, _auto_col_widths(rows, 8.5, CW))
            elif sh.get("empty_note"):
                put(sh["empty_note"], size=9, color=_F.MUTED, gap_after=4)

    # ── Summary ───────────────────────────────────────────────────────────────
    summary_raw = company.get("summary", "")
    if summary_raw:
        section_heading("公司簡介")
        for heading, body in _md_sections(summary_raw):
            if heading:
                summary_h3(heading)
            pdf_summary_body(body.split("\n"))

    # ── 專利 ───────────────────────────────────────────────────────────────────
    patents = company.get("patents") or []
    if patents:
        section_heading("專利")
        put(f"共 {len(patents)} 筆（更新：{patents[0].get('fetched_at', '')}）",
            size=9, color=_F.MUTED, gap_after=4)
        prows = _patent_table_rows(patents)
        pdf_table(prows, _auto_col_widths(prows, 8.5, CW))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
