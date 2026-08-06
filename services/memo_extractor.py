"""
Extract Call Memo fields from a meeting transcript using Claude CLI.
Returns a dict matching the 24-field template schema (+ interview_date on extraction).
"""
import asyncio
import json
import re
from datetime import datetime
from pathlib import Path

from services import claude_client

# Ordered field definitions: key -> (label, short_description)
FIELDS: list[tuple[str, str, str]] = [
    ("deal_source",        "案件來源",                   "例：自行開發、某人介紹；若未提及請留空"),
    ("interviewees",       "受訪人",                     "受訪者姓名與職稱，多人以頓號分隔"),
    ("paid_in_capital",    "實收資本額",                  "NT$ 金額，例：5,000萬"),
    ("address",            "地址",                       "公司登記地址或廠址"),
    ("founding_date",      "設立日期",                   "例：2018年 或 2018/03/01"),
    ("underwriter",        "承銷商",                     "輔導券商名稱，若未提及填空"),
    ("auditor",            "會計師事務所",                "簽證會計師事務所，若未提及填空"),
    ("chairman",           "董事長",                     "董事長姓名"),
    ("general_manager",    "總經理",                     "總經理姓名"),
    ("headcount",          "員工人數",                   "數字，例：120人"),
    ("ipo_timeline",       "公開發行及上市櫃時程/募資規劃", "IPO 目標年份、目前募資輪次等時間性資訊"),
    ("investment_terms",   "增資計畫或投資條件",          "本次募資總額、釋出股比、預計 close 時程"),
    ("business_revenue",   "主要業務、產品營收比重",       "核心業務說明及各產品/服務的營收佔比"),
    ("financials",         "財務狀況",                   "近期營收、淨利、年增率等財務數據"),
    ("management_team",    "經營團隊背景",                "創辦人/CEO、CTO、CFO 的背景與經歷"),
    ("board_shareholding", "董監或主要股東持股情形",       "主要股東名稱與持股比例"),
    ("recent_development", "公司發展近況",                "近期重大里程碑、產品進展、合作案"),
    ("major_customers",    "主要銷貨客戶",                "前幾大客戶名稱與佔比"),
    ("major_suppliers",    "主要進貨廠商",                "主要原物料或零件供應商"),
    ("factory_capacity",   "廠房及產能使用情形",          "廠房地點、產能規模、目前使用率"),
    ("competitors",        "國內外主要競爭對手",          "直接競爭者名稱及差異化分析"),
    ("industry_trends",    "產業發展趨勢",                "產業現況、市場規模、未來展望"),
    ("risk_tracking",      "風險評估及追蹤事項",          "主要風險點與需持續追蹤的議題"),
    ("conclusion",         "評估結論與建議",              "訪談整體評估與後續建議行動"),
]

FIELD_KEYS = [f[0] for f in FIELDS]

# 抽取時額外請 AI 從逐字稿判斷訪談日期（不併進 FIELDS——FIELDS 是 memo 正文 24 欄的
# 唯一真理來源，供 MemoSave / DOCX 範本 / serialize_memo 共用；interview_date 另走流程）。
_EXTRACT_FIELDS = [("interview_date", "訪談日期", "逐字稿中提到的訪談/會議日期，格式 YYYY/MM/DD；未提及請留空")] + FIELDS
_EXTRACT_KEYS = [f[0] for f in _EXTRACT_FIELDS]

_CHUNK_CHARS = 12000
_CHUNK_OVERLAP = 500
PROMPT_VERSION = "evidence-v6.1-synthesized"

# Keep each synthesis prompt focused enough that the model can actually edit
# prose instead of falling back to copying a long evidence dump.  Evidence is
# still stored in full in the audit file; these groups only control the final
# presentation pass.
_SYNTHESIS_GROUPS: tuple[tuple[str, ...], ...] = (
    (
        "deal_source", "interviewees", "paid_in_capital", "address",
        "founding_date", "underwriter", "auditor", "chairman",
        "general_manager", "headcount",
    ),
    (
        "ipo_timeline", "investment_terms", "financials",
        "board_shareholding",
    ),
    (
        "business_revenue", "recent_development", "major_customers",
        "major_suppliers", "factory_capacity",
    ),
    (
        "management_team", "competitors", "industry_trends",
        "risk_tracking", "conclusion",
    ),
)

_QUESTION_CUES = re.compile(r"(?:有沒有|有没有|是不是|會不會|会不会|你覺得|你觉得|嗎|吗|[?？])(?:[。.!！]?)$")
_THIRD_PARTY_EXAMPLE_CUES = re.compile(r"(?:比如|例如|假設|假设|uber|wechat|特斯拉)", re.I)
_FIRST_PARTY_CUES = re.compile(r"(?:我們|我们|本公司|我司|本團隊|本团队)")

# Generic investment-signal grammar. No company-, year-, amount-, or country-
# specific literals belong here.
_GENERIC_SIGNAL_RULES: tuple[tuple[str, re.Pattern], ...] = (
    ("financials", re.compile(
        r"(?:nt\$|us\$|rmb|jpy|usd|twd|新台幣|人民幣|美元|日圓|日元)\s*"
        r"(?:\d[\d,.]*|[一二三四五六七八九十百千兩两]+)(?:萬|万|億|亿|\s*(?:k|m|million|billion))?|"
        r"(?:\d[\d,.]*|[一二三四五六七八九十百千兩两]+)"
        r"(?:(?:萬|万|億|亿)(?:元|塊|美元|日圓|日元)?|(?:元|塊|美元|日圓|日元)|\s*(?:k|m|million|billion)\b)|"
        r"(?:\d+(?:\.\d+)?\s*%|百分之[一-鿿\d.]+)|"
        r"(?:營收|营收|淨利|净利|毛利|獲利|获利|成長率|成长率|複合成長|复合成长)")),
    ("headcount", re.compile(r"(?:\d[\d,]*|[一二三四五六七八九十百千兩两]+)\s*(?:個|个)?(?:人|名)(?:員工|员工|團隊|团队|工程師|工程师|研發|研发|rd)?", re.I)),
    ("recent_development", re.compile(r"(?:19|20)\d{2}(?:年)?|\bhr\b|行政|財務|财务|組織調整|组织调整|里程碑", re.I)),
    ("major_customers", re.compile(r"客戶|客户|市場|海外|出口|跨境|國際|国际|全球")),
    ("management_team", re.compile(r"核心能力|關鍵人員|关键人员|keyman|創辦人|创办人|執行長|执行长|ceo|cto|cfo", re.I)),
    ("risk_tracking", re.compile(r"風險|风险|商機|商机|下降|流失|離職|离职|法規|現金流|现金流|中斷|取代|不承接|不接")),
    ("factory_capacity", re.compile(r"產能|产能|良率|廠房|厂房|稼動率|稼动率|月產|月产")),
    ("ipo_timeline", re.compile(r"\bipo\b|公開發行|公开发行|上市|上櫃|上柜|興櫃|兴柜", re.I)),
    ("investment_terms", re.compile(r"投資人|投资人|募資|募资|增資|增资|估值|併購|并购|股權|股权|表決|表决|釋股|释股|close", re.I)),
)


def prepare_transcript(transcript: str, source_filename: str = "") -> str:
    """Remove generated Markdown summaries/metadata when a raw transcript section exists."""
    if Path(source_filename).suffix.lower() != ".md":
        return transcript
    marker = re.search(r"(?im)^##\s*逐字稿\s*$", transcript)
    return transcript[marker.end():].lstrip() if marker else transcript


def infer_date_from_filename(source_filename: str) -> str:
    """Infer YYYY/MM/DD from an explicit date token in a filename."""
    stem = Path(source_filename).stem
    for pattern in (
        r"(?<!\d)(20\d{2})(\d{2})(\d{2})(?!\d)",
        r"(?<!\d)(20\d{2})[-_.](\d{1,2})[-_.](\d{1,2})(?!\d)",
    ):
        match = re.search(pattern, stem)
        if not match:
            continue
        candidate = "/".join(match.groups())
        try:
            return datetime.strptime(candidate, "%Y/%m/%d").strftime("%Y/%m/%d")
        except ValueError:
            continue
    return ""


def _split_transcript(transcript: str) -> list[str]:
    """超長逐字稿切成有重疊的片段，避免在句子中間切斷漏抽。"""
    if len(transcript) <= _CHUNK_CHARS:
        return [transcript]
    chunks, start, step = [], 0, _CHUNK_CHARS - _CHUNK_OVERLAP
    while start < len(transcript):
        chunks.append(transcript[start:start + _CHUNK_CHARS])
        start += step
    return chunks


async def extract_from_transcript(
    company_name: str,
    transcript: str,
    engine: str = "",
    source_filename: str = "",
) -> dict:
    fields, _audit = await extract_with_audit(
        company_name, transcript, engine=engine, source_filename=source_filename
    )
    return fields


async def extract_with_audit(
    company_name: str,
    transcript: str,
    engine: str = "",
    source_filename: str = "",
) -> tuple[dict, dict]:
    """
    從逐字稿抽取所有 Call Memo 欄位（+ interview_date）。
    逐字稿超長時分段抽取「事實 + 原文證據」，累積去重後再統整成 24 欄。
    AI 回不出可解析 JSON 時 raise ValueError（不偽裝成 24 欄全空的『成功』）。
    """
    transcript = prepare_transcript(transcript, source_filename)
    chunks = _split_transcript(transcript)
    evidence: dict[str, list[dict[str, str]]] = {k: [] for k in _EXTRACT_KEYS}
    chunk_results = await asyncio.gather(*(
        _extract_chunk_dual(company_name, chunk, engine) for chunk in chunks
    ))
    candidate_chunks = _split_transcript(_material_candidate_text(transcript))
    candidate_results = await asyncio.gather(*(
        _extract_material_facts_once(company_name, chunk, engine)
        for chunk in candidate_chunks if chunk.strip()
    ))
    chunk_results.extend([[result] for result in candidate_results])
    deterministic = _deterministic_material_evidence(transcript)
    for items in deterministic.values():
        for item in items:
            item["source"] = "deterministic_safety_net"
    chunk_results.append([deterministic])
    for parts in chunk_results:
        for part in parts:
            for key, items in part.items():
                seen = {(e["fact"], e["quote"], e["timestamp"]) for e in evidence[key]}
                for item in items:
                    anchor = _normalized_text(str(item.get("anchor_quote") or ""))
                    if anchor and any(
                        anchor in _normalized_text(existing["quote"])
                        for existing in evidence[key]
                    ):
                        continue
                    identity = (item["fact"], item["quote"], item["timestamp"])
                    if identity not in seen:
                        evidence[key].append(item)
                        seen.add(identity)

    evidence_by_id: dict[str, list[dict[str, str]]] = {}
    for key in _EXTRACT_KEYS:
        evidence_by_id[key] = [
            {**item, "id": f"{key}:{index}"}
            for index, item in enumerate(evidence[key], start=1)
        ]

    fields, coverage = await _synthesize_fields_from_evidence(
        company_name, evidence_by_id, engine
    )
    if not fields["interview_date"]:
        fields["interview_date"] = infer_date_from_filename(source_filename)
    return fields, {
        "evidence": evidence_by_id,
        "coverage": coverage,
    }


async def _extract_chunk_dual(
    company_name: str, transcript: str, engine: str
) -> list[dict]:
    """Run complementary extractors; one timeout must not discard the other."""
    results = await asyncio.gather(
        _extract_evidence_once(company_name, transcript, engine),
        _extract_material_facts_once(company_name, transcript, engine),
        return_exceptions=True,
    )
    successful = [result for result in results if isinstance(result, dict)]
    if successful:
        return successful
    first_error = next((result for result in results if isinstance(result, Exception)), None)
    raise ValueError(f"AI 無法完成逐字稿分段抽取：{first_error}")


def _json_object(raw: str) -> dict:
    raw = re.sub(r"^```[a-z]*\n?", "", raw.strip(), flags=re.MULTILINE)
    raw = re.sub(r"\n?```$", "", raw.strip(), flags=re.MULTILINE)
    match = re.search(r"\{[\s\S]*\}", raw)
    if not match:
        raise ValueError("AI 未回傳可解析的 call memo（請重試或換引擎）")
    try:
        data = json.loads(match.group())
    except Exception as e:
        raise ValueError("AI 回傳的 call memo 非合法 JSON（請重試或換引擎）") from e
    if not isinstance(data, dict):
        raise ValueError("AI 回傳的 call memo 不是 JSON 物件（請重試或換引擎）")
    return data


def _normalized_text(value: str) -> str:
    return re.sub(r"\s+", "", value or "")


def _material_candidate_text(transcript: str) -> str:
    """Select high-signal transcript windows for a focused omission audit."""
    lines = transcript.splitlines()
    selected: set[int] = set()
    for index, line in enumerate(lines):
        content = re.sub(r"^\s*\[\d{2}:\d{2}:\d{2}\]\s*", "", line).lower()
        if any(pattern.search(content) for _, pattern in _GENERIC_SIGNAL_RULES):
            selected.update(range(max(0, index - 2), min(len(lines), index + 3)))
    return "\n".join(lines[index] for index in sorted(selected))


def _deterministic_material_evidence(transcript: str) -> dict:
    """Last-resort verbatim capture for facts models commonly omit."""
    lines = transcript.splitlines()
    result: dict[str, list[dict[str, str]]] = {key: [] for key in _EXTRACT_KEYS}
    seen: set[tuple[str, int, int]] = set()
    for index, line in enumerate(lines):
        content = re.sub(r"^\s*\[\d{2}:\d{2}:\d{2}\]\s*", "", line)
        for field, pattern in _GENERIC_SIGNAL_RULES:
            if not pattern.search(content):
                continue
            window_text = "\n".join(lines[max(0, index - 2):min(len(lines), index + 3)])
            # Questions and unrelated examples are candidates for AI review, but
            # are unsafe to force directly into the final memo.
            if _QUESTION_CUES.search(content):
                continue
            if _THIRD_PARTY_EXAMPLE_CUES.search(window_text) and not _FIRST_PARTY_CUES.search(window_text):
                continue
            start, end = max(0, index - 2), min(len(lines), index + 3)
            identity = (field, start, end)
            if identity in seen:
                continue
            seen.add(identity)
            quote = "\n".join(lines[start:end]).strip()
            # Keep surrounding lines as audit evidence, but only render the
            # matched assertion as the fact. This avoids turning context and
            # interviewer chatter into memo prose.
            fact = content.strip()
            timestamp_match = re.search(r"\[(\d{2}:\d{2}:\d{2})\]", line)
            result[field].append({
                "fact": fact,
                "quote": quote,
                "timestamp": timestamp_match.group(1) if timestamp_match else "",
                "anchor_quote": line.strip(),
            })
    return result


async def _extract_evidence_once(company_name: str, transcript: str, engine: str) -> dict:
    fields_desc = "\n".join(
        f'  "{key}": "{label}（{desc}）"'
        for key, label, desc in _EXTRACT_FIELDS
    )

    prompt = f"""你是逐字稿事實抽取器。以下是與「{company_name}」的訪談逐字稿其中一段。

請為每個欄位找出所有明確事實與逐字證據，以 JSON 回傳：
- 每項格式為 {{"fact":"繁體中文事實", "quote":"原文逐字引用", "timestamp":"HH:MM:SS"}}。
- quote 必須逐字複製本段內容，不得翻譯、改寫或修正簡繁體；找不到原文引用就不要輸出。
- fact 只能表達 quote 直接支持的內容，不得加入常識、推測、評價或外部資訊。
- 同一事實可放入多個真正相關的欄位，以免後續漏掉重要內容。
- 這是高召回率抽取，不是摘要；寧可多列也不可只挑幾個重點。
- 逐項檢查：金額與比例、年份與時程、人數、客戶地區與類型、產品與收入模式、核心能力、經營原則、明確不做的事、風險、轉折點、海外布局、募資與 IPO，有提及都要列入。
- 本段未提及的欄位回傳空陣列 []。
- generated_at、檔案建立時間及逐字稿產生時間不是訪談日期。
- 回傳純 JSON，不要加 markdown code block 或說明。

需提取的欄位：
{{
{fields_desc}
}}

逐字稿內容：
---
{transcript}
---

請直接回傳 JSON 物件，每個 key 的 value 都是陣列。"""

    raw = await asyncio.to_thread(claude_client.ask, prompt, 180, None, engine)
    data = _json_object(raw)
    return _validated_evidence(data, transcript)


async def _extract_material_facts_once(
    company_name: str, transcript: str, engine: str
) -> dict:
    field_choices = "\n".join(
        f'- {key}: {label}' for key, label, _ in _EXTRACT_FIELDS
    )
    prompt = f"""你是投資訪談的高召回率事實稽核員。請按原文出現順序，掃描「{company_name}」這段逐字稿的每個投資重要事實。

這不是摘要。特別不可漏掉：
- 金額、比例、人數、年份、時程與成長率
- 客戶地區與類型、訂單變化、收入模式與海外布局
- 核心能力、組織變化、關鍵人員、經營原則與明確不做的事
- 疫情、AI、人才、客戶、競爭、法規、現金流等風險
- 募資、併購、股權、表決權、IPO 動機與資金用途

每筆事實指定一個最適合的 field，格式：
{{"facts":[{{"field":"field_key","fact":"繁體中文事實","quote":"原文逐字引用","timestamp":"HH:MM:SS"}}]}}

field 只能從以下選擇：
{field_choices}

quote 必須逐字存在於本段原文；無法提供引用就不得輸出。只回傳純 JSON。

逐字稿：
---
{transcript}
---
"""
    raw = await asyncio.to_thread(claude_client.ask, prompt, 180, None, engine)
    data = _json_object(raw)
    grouped: dict[str, list[dict]] = {key: [] for key in _EXTRACT_KEYS}
    for item in data.get("facts", []):
        if isinstance(item, dict) and item.get("field") in grouped:
            grouped[item["field"]].append(item)
    return _validated_evidence(grouped, transcript)


def _validated_evidence(data: dict, transcript: str) -> dict:
    normalized_chunk = _normalized_text(transcript)
    result: dict[str, list[dict[str, str]]] = {k: [] for k in _EXTRACT_KEYS}
    for key in _EXTRACT_KEYS:
        items = data.get(key, [])
        if not isinstance(items, list):
            continue
        for item in items:
            if not isinstance(item, dict):
                continue
            fact = str(item.get("fact") or "").strip()
            quote = str(item.get("quote") or "").strip()
            timestamp = str(item.get("timestamp") or "").strip()
            # Hard hallucination gate: the cited quote must exist verbatim in this chunk.
            if fact and quote and _normalized_text(quote) in normalized_chunk:
                result[key].append({"fact": fact, "quote": quote, "timestamp": timestamp})
    return result


def _deduplicated_facts(items: list[dict[str, str]]) -> list[dict[str, str]]:
    """Select readable evidence and remove exact/containment duplicates.

    Verbatim safety-net captures are always audit-only. They help reviewers find
    possible omissions, but lack the semantic judgment required for publication
    (for example, a third-party factory example can mention "產能").
    """
    selected: list[dict[str, str]] = []
    normalized_facts: list[str] = []
    model_items = [
        item for item in items
        if item.get("source") != "deterministic_safety_net"
    ]
    for item in model_items:
        fact = str(item.get("fact") or "").strip()
        normalized = _normalized_text(fact)
        if not normalized:
            continue
        contained_at = next(
            (i for i, existing in enumerate(normalized_facts)
             if normalized in existing or existing in normalized),
            None,
        )
        if contained_at is None:
            selected.append(item)
            normalized_facts.append(normalized)
        elif len(normalized) > len(normalized_facts[contained_at]):
            selected[contained_at] = item
            normalized_facts[contained_at] = normalized
    return selected


def _fallback_field_text(items: list[dict[str, str]], limit: int = 8) -> str:
    """Readable non-AI fallback: no duplicated sentence terminators or raw flood."""
    facts = []
    for item in _deduplicated_facts(items):
        fact = str(item.get("fact") or "").strip().rstrip("。；; ")
        if fact:
            facts.append(fact)
        if len(facts) >= limit:
            break
    return ("；".join(facts) + "。") if facts else ""


def _interview_date_from_evidence(items: list[dict[str, str]]) -> str:
    for item in items:
        fact = str(item.get("fact") or "")
        match = re.search(r"\b(20\d{2})[/-](\d{1,2})[/-](\d{1,2})\b", fact)
        if not match:
            continue
        try:
            return datetime.strptime(
                "/".join(match.groups()), "%Y/%m/%d"
            ).strftime("%Y/%m/%d")
        except ValueError:
            continue
    return ""


async def _synthesize_fields_from_evidence(
    company_name: str,
    evidence: dict[str, list[dict[str, str]]],
    engine: str,
) -> tuple[dict[str, str], dict[str, dict]]:
    """Turn verified evidence into concise, readable Call Memo prose.

    The model may select and merge evidence, but every non-empty field must cite
    at least one valid evidence ID. Raw evidence remains in the audit payload.
    """
    fields = {key: "" for key in _EXTRACT_KEYS}
    coverage: dict[str, dict] = {}
    fields["interview_date"] = _interview_date_from_evidence(
        evidence["interview_date"]
    )
    coverage["interview_date"] = {
        "evidence_count": len(evidence["interview_date"]),
        "used_count": 1 if fields["interview_date"] else 0,
        "synthesized": False,
    }

    results = await asyncio.gather(*(
        _synthesize_field_group(company_name, group, evidence, engine)
        for group in _SYNTHESIS_GROUPS
    ), return_exceptions=True)

    for group, result in zip(_SYNTHESIS_GROUPS, results):
        parsed = result if isinstance(result, dict) else {}
        for key in group:
            items = evidence[key]
            valid_ids = {item["id"] for item in items}
            entry = parsed.get(key) if isinstance(parsed, dict) else None
            text = ""
            used_ids: list[str] = []
            if isinstance(entry, dict):
                candidate = str(entry.get("text") or "").strip()
                cited = entry.get("evidence_ids") or []
                if isinstance(cited, list):
                    used_ids = [str(i) for i in cited if str(i) in valid_ids]
                # A non-empty answer without a real evidence citation is not
                # trusted. Fall back to bounded deterministic prose instead.
                if candidate and used_ids:
                    text = _normalize_memo_prose(candidate)
            if not text and items:
                text = _fallback_field_text(items)
            fields[key] = text
            coverage[key] = {
                "evidence_count": len(items),
                "used_count": len(set(used_ids)),
                "synthesized": bool(text and used_ids),
            }
    return fields, coverage


async def _synthesize_field_group(
    company_name: str,
    keys: tuple[str, ...],
    evidence: dict[str, list[dict[str, str]]],
    engine: str,
) -> dict:
    payload: dict[str, list[dict[str, str]]] = {}
    for key in keys:
        payload[key] = [
            {"id": item["id"], "fact": item["fact"]}
            for item in _deduplicated_facts(evidence[key])
        ]

    prompt = f"""你是台灣私募股權團隊的 Call Memo 編輯。請將「{company_name}」的已驗證事實整理成可直接放進正式備忘錄的欄位文字。

規則：
1. 只能使用下方 evidence，不得加入外部資訊、推測或常識。
2. 合併語意重複的事實；不要逐條照抄，不要暴露逐字稿碎片。
3. 全文使用自然、專業的台灣繁體中文；修正明顯語音辨識用字，但不可改變事實。
4. 短欄位（姓名、日期、地址、金額）只填直接答案。敘述欄位通常 100–350 字，資訊很多時最多 500 字。
5. 使用完整句子與正常標點。禁止「。；」「；。」，禁止同一句換句話說重複出現。
6. 未有 evidence 的欄位輸出空字串。
7. 每個非空欄位列出實際採用的 evidence_ids；ID 必須來自該欄位。
8. 只輸出 JSON，格式如下：
{{"field_key":{{"text":"整理後文字","evidence_ids":["field_key:1"]}}}}

欄位邊界：
- financials 只放公司已發生的營收、獲利、成長率或明確財務狀況；一般營運風險放 risk_tracking。
- factory_capacity 只放公司本身的廠房、實體產能、良率或稼動率；軟體交付人力、開發速度及第三方案例均留空。
- competitors 只整理實際競爭者、替代方案或受訪者明確描述的競爭情境，不要拿公司管理優勢填充。

已驗證 evidence：
{json.dumps(payload, ensure_ascii=False)}
"""
    raw = await asyncio.to_thread(
        claude_client.ask, prompt, 240, None, engine
    )
    data = _json_object(raw)
    return {key: data.get(key, {}) for key in keys}


def _normalize_memo_prose(value: str) -> str:
    text = re.sub(r"[ \t]+", " ", value or "").strip()
    text = re.sub(r"。\s*；", "。", text)
    text = re.sub(r"；\s*。", "。", text)
    text = re.sub(r"；{2,}", "；", text)
    text = re.sub(r"。{2,}", "。", text)
    return text


def fill_template(company: dict, memo: dict, interview_date: str = "") -> bytes:
    """Fill the Call Memo .docx template with memo fields, return bytes."""
    from docx import Document
    from docx.oxml.ns import qn
    from copy import deepcopy
    import io

    template_path = Path(__file__).parent.parent / "data" / "call_memo_template.docx"
    doc = Document(str(template_path))

    # ── Fill header paragraphs (訪談日期 / 評估人) ──────────────────────────
    for para in doc.paragraphs:
        if "訪談日期：" in para.text and interview_date:
            for run in para.runs:
                if "2025/X/X" in run.text or "X/X" in run.text:
                    run.text = run.text.replace("2025/X/X", interview_date).replace("X/X", interview_date)
                    break

    # ── Label → field key mapping（單一來源 FIELDS 衍生，不再手抄 24 條）─────────
    LABEL_MAP = {label: key for key, label, _ in FIELDS}
    LABEL_MAP.update({f"{label}：": key for key, label, _ in FIELDS})  # 範本用「：」結尾
    LABEL_MAP.update({   # 範本文字與 FIELDS 標籤不同的變體
        "公司名稱：": "_company_name",
        "會計師：": "auditor",
        "董監(或主要股東)持股情形：": "board_shareholding",
    })

    company_name = company.get("name", "")

    def _get_value(field_key: str) -> str:
        if field_key == "_company_name":
            return company_name
        return memo.get(field_key, "")

    def _fill_cell(cell, value: str):
        """Keep first paragraph (label), remove rest, add value paragraphs."""
        paragraphs = cell.paragraphs
        if not paragraphs:
            return

        # Remove all paragraphs after the first
        tc = cell._tc
        for p in paragraphs[1:]:
            tc.remove(p._p)

        # First paragraph: keep label runs, remove any non-bold value runs
        first_p = cell.paragraphs[0]
        # Find where label ends (last bold run ending with ：)
        label_end_idx = -1
        for i, run in enumerate(first_p.runs):
            if run.bold or (run.text and run.text.strip().endswith("：")):
                label_end_idx = i

        # Remove runs after label
        for run in first_p.runs[label_end_idx + 1:]:
            first_p._p.remove(run._r)

        if not value:
            return

        # Add value: same paragraph for short values, new paragraphs for multiline
        lines = [l for l in value.split("\n") if l.strip()]
        if not lines:
            return

        # Add first line to label paragraph
        r = first_p.add_run(" " + lines[0])
        r.bold = False

        # Additional lines as new paragraphs
        for line in lines[1:]:
            new_p = cell.add_paragraph()
            new_p.add_run(line)

    # ── Iterate table cells ───────────────────────────────────────────────────
    if doc.tables:
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                matched_key = None
                matched_label = None
                for label, key in LABEL_MAP.items():
                    if cell_text.startswith(label) or cell_text == label.rstrip("：") + "：":
                        matched_key = key
                        matched_label = label
                        break
                if matched_key:
                    _fill_cell(cell, _get_value(matched_key))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
